import docx
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import zipfile
import re
import os
import base64
import requests

template_path = r"C:\Users\Dell\Desktop\Fyp\FYP_Product-Based_Template CS.docx"
original_doc_path = r"C:\Users\Dell\Desktop\Fyp\BIZIONARY ERP - Project Documentation.docx"
md_path = r"C:\Users\Dell\Desktop\Fyp\PROJECT_DOCUMENTATION.md"
output_path = r"C:\Users\Dell\Desktop\Fyp\BIZIONARY ERP - Project Documentation (Reformatted).docx"

# Temporary directory for image extraction and rendering
scratch_dir = r"C:\Users\Dell\.gemini\antigravity-ide\brain\b1a9658c-28bd-4dd4-ad82-d744277308e4\scratch"
os.makedirs(scratch_dir, exist_ok=True)

# Helper function to add background color to paragraphs (for code blocks)
def set_paragraph_shading(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    pPr.append(shd)

# Helper function to insert Table of Contents and lists fields in Word
def add_field_code(paragraph, field_text):
    p = paragraph._p
    
    # 1. Field Begin
    r1 = OxmlElement('w:r')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    r1.append(fldChar1)
    p.append(r1)
    
    # 2. Instruction Text
    r2 = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_text
    r2.append(instrText)
    p.append(r2)
    
    # 3. Separate Field
    r3 = OxmlElement('w:r')
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    r3.append(fldChar2)
    p.append(r3)
    
    # 4. Result area placeholder
    r_placeholder = OxmlElement('w:r')
    p.append(r_placeholder)
    
    # 5. Field End
    r4 = OxmlElement('w:r')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r4.append(fldChar3)
    p.append(r4)

# Tokenize and parse inline markdown formatting (bold, italic, inline code, links)
def add_markdown_runs(paragraph, text):
    pattern = r'(\*\*.*?\*\*|__.*?__|`.*?`|_[^_]+_|\*.*?\*|\[.*?\]\(.*?\))'
    tokens = re.split(pattern, text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith('__') and token.endswith('__'):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith('`') and token.endswith('`'):
            r = paragraph.add_run(token[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(163, 21, 21) # dark red for code text
        elif token.startswith('_') and token.endswith('_'):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        elif token.startswith('*') and token.endswith('*'):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        elif token.startswith('[') and ']' in token and '(' in token and token.endswith(')'):
            match = re.match(r'\[(.*?)\]\((.*?)\)', token)
            if match:
                link_text = match.group(1)
                r = paragraph.add_run(link_text)
                r.font.underline = True
                r.font.color.rgb = RGBColor(9, 105, 217) # professional blue links
        else:
            paragraph.add_run(token)

def clear_body_after(doc, paragraph_index):
    p_elem = doc.paragraphs[paragraph_index]._element
    body = doc.element.body
    children = list(body)
    idx = children.index(p_elem)
    for child in children[idx:]:
        if child.tag.endswith('sectPr'):
            continue
        body.remove(child)

def get_mermaid_image(mermaid_code, filename):
    try:
        graph_bytes = mermaid_code.encode("utf8")
        base64_bytes = base64.urlsafe_b64encode(graph_bytes)
        base64_string = base64_bytes.decode("ascii")
        url = f"https://mermaid.ink/img/{base64_string}"
        
        print(f"  Attempting to render: {filename} via mermaid.ink...")
        r = requests.get(url, timeout=20)
        if r.status_code == 200:
            img_path = os.path.join(scratch_dir, filename)
            with open(img_path, 'wb') as f:
                f.write(r.content)
            print(f"    Success rendering {filename}")
            return img_path
        else:
            print(f"    Failed rendering {filename} (HTTP {r.status_code})")
    except Exception as e:
        print(f"    Error rendering {filename}: {e}")
    return None

def main():
    print("Step 1: Extracting images from original docx...")
    if os.path.exists(original_doc_path):
        with zipfile.ZipFile(original_doc_path) as z:
            for img_name in ['image6.png', 'image7.png', 'image47.png']:
                zip_path = f"word/media/{img_name}"
                if zip_path in z.namelist():
                    data = z.read(zip_path)
                    with open(os.path.join(scratch_dir, img_name), 'wb') as f:
                        f.write(data)
                    print(f"  Extracted {img_name}")
                else:
                    print(f"  Warning: {img_name} not found in zip container")

    print("\nStep 2: Loading template docx...")
    doc = docx.Document(template_path)
    
    print("Step 3: Modifying cover page details...")
    p_name = doc.paragraphs[3]
    p_name.text = ""
    r_name = p_name.add_run("BIZIONARY ERP SYSTEM")
    r_name.bold = True
    r_name.font.name = 'Book Antiqua'
    r_name.font.size = Pt(22)
    p_name.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    p_sub = doc.paragraphs[4]
    p_sub.text = ""
    r_sub = p_sub.add_run("Final Year Project Technical Documentation")
    r_sub.bold = True
    r_sub.font.name = 'Book Antiqua'
    r_sub.font.size = Pt(14)
    r_sub.font.italic = True
    p_sub.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    p_advisor = doc.paragraphs[5]
    p_advisor.text = ""
    r_adv = p_advisor.add_run("Project Advisor:\n[Advisor Name]")
    r_adv.bold = True
    r_adv.font.name = 'Book Antiqua'
    r_adv.font.size = Pt(12)
    p_advisor.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    p_submitted = doc.paragraphs[7]
    p_submitted.text = ""
    r_subm = p_submitted.add_run("Submitted By:\n[Student Name / Roll Number]")
    r_subm.bold = True
    r_subm.font.name = 'Book Antiqua'
    r_subm.font.size = Pt(12)
    p_submitted.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    p_session = doc.paragraphs[10]
    p_session.text = ""
    r_ses = p_session.add_run("Session: [Session Years]")
    r_ses.bold = True
    r_ses.font.name = 'Book Antiqua'
    r_ses.font.size = Pt(12)
    p_session.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    print("Step 4: Setting structural pages placeholders...")
    # Dedication
    doc.paragraphs[18].text = ""
    r_ded = doc.paragraphs[18].add_run("[Dedication content goes here...]")
    r_ded.font.name = 'Book Antiqua'
    r_ded.font.size = Pt(11)
    r_ded.font.italic = True
    doc.paragraphs[18].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    # Acknowledgment
    doc.paragraphs[75].text = ""
    r_ack = doc.paragraphs[75].add_run("[Acknowledgment content goes here...]")
    r_ack.font.name = 'Book Antiqua'
    r_ack.font.size = Pt(11)

    # Project metadata fields
    metadata_fields = {
        87: "Project Title\t\t\tBIZIONARY ERP SYSTEM",
        88: "Objective\t\t\tEnterprise Resource Planning (ERP) & Business Intelligence platform",
        90: "Undertaken by\t\t[Student Names / Roll Numbers]",
        92: "Supervised by\t\t[Supervisor Name]",
        93: "Starting Date\t\t\t[Start Date]",
        95: "Completion Date\t\t[Completion Date]",
        97: "Tools Used\t\t\tDjango REST Framework, React, PostgreSQL/SQLite, Groq API",
        98: "Operating System\t\tCross-platform (Windows, Linux, macOS)",
        99: "Documentation\t\tTechnical Documentation"
    }
    for idx, text in metadata_fields.items():
        doc.paragraphs[idx].text = ""
        r_m = doc.paragraphs[idx].add_run(text)
        r_m.font.name = 'Book Antiqua'
        r_m.font.size = Pt(11)

    # Plagiarism Report & Abstract placeholders
    doc.paragraphs[109].text = "[Plagiarism Report goes here...]"
    doc.paragraphs[109].runs[0].font.name = 'Book Antiqua'
    doc.paragraphs[109].runs[0].font.size = Pt(11)
    doc.paragraphs[110].text = "[Abstract content goes here...]"
    doc.paragraphs[110].runs[0].font.name = 'Book Antiqua'
    doc.paragraphs[110].runs[0].font.size = Pt(11)

    print("Step 5: Clearing template chapters and outline...")
    clear_body_after(doc, 113)

    print("Step 5.5: Reading markdown headings for pre-populated TOC...")
    headings = []
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    md_lines = md_text.split('\n')
    for line in md_lines:
        striped = line.strip()
        if striped.startswith("## "):
            headings.append((1, striped[3:].strip()))
        elif striped.startswith("### "):
            headings.append((2, striped[4:].strip()))
        elif striped.startswith("#### "):
            headings.append((3, striped[5:].strip()))

    figures = [
        "Figure 1: High-Level Client-Server Architecture Diagram",
        "Figure 2: Dynamic Sales Excel Ingestion Sequence Flow",
        "Figure 3: Agentic Chatbot RAG Query Resolution Process",
        "Figure 4: User Authentication and RBAC Flow",
        "Figure 5: Double-Entry Ledger Transaction Posting Flow",
        "Figure 6: Supplier Procurement Ordered Slip Lifecycle",
        "Figure 7: Entity Relationship Diagram (ERD) Schema",
        "Figure 8: Executive Dashboard Overview",
        "Figure 9: Dashboard Performance Insights",
        "Figure 10: Accounts and Finance Dashboard",
        "Figure 11: Products Catalog Grid",
        "Figure 12: Stock Management Control Center",
        "Figure 13: Warehouse Stock Breakdown Modal",
        "Figure 14: Pending Incoming Stock Modal",
        "Figure 15: Sales Performance Dashboard",
        "Figure 16: Sales Transaction Log Table",
        "Figure 17: Ordered Slips Dashboard",
        "Figure 18: AI Chatbot Interface"
    ]

    print("Step 6: Creating Table of Contents and Lists of Figures/Tables fields (Pre-populated)...")
    # Table of Contents
    p_toc_title = doc.add_paragraph()
    r_toc_t = p_toc_title.add_run("Table of Contents")
    r_toc_t.bold = True
    r_toc_t.font.name = 'Book Antiqua'
    r_toc_t.font.size = Pt(16)
    p_toc_title.style = 'Heading 1'
    p_toc_title.paragraph_format.space_before = Pt(18)
    p_toc_title.paragraph_format.space_after = Pt(12)
    
    p_toc_inst = doc.add_paragraph()
    r_toc_inst = p_toc_inst.add_run("Select this area, right-click, and select 'Update Field' to refresh page numbers.")
    r_toc_inst.font.italic = True
    r_toc_inst.font.name = 'Book Antiqua'
    r_toc_inst.font.size = Pt(10)
    
    p_toc_begin = doc.add_paragraph()
    r1 = OxmlElement('w:r')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    r1.append(fldChar1)
    p_toc_begin._element.append(r1)
    
    r2 = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    r2.append(instrText)
    p_toc_begin._element.append(r2)
    
    r3 = OxmlElement('w:r')
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    r3.append(fldChar2)
    p_toc_begin._element.append(r3)
    
    # Pre-populate TOC
    for level, h_text in headings:
        p_entry = doc.add_paragraph(style=f'toc {level}')
        p_entry.paragraph_format.space_before = Pt(0)
        p_entry.paragraph_format.space_after = Pt(2)
        r_txt = p_entry.add_run(h_text)
        r_txt.font.name = 'Book Antiqua'
        r_txt.font.size = Pt(10)
        p_entry.add_run('\t')
        r_pg = p_entry.add_run('1')
        r_pg.font.name = 'Book Antiqua'
        r_pg.font.size = Pt(10)
        
    p_toc_end = doc.add_paragraph()
    r4 = OxmlElement('w:r')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r4.append(fldChar3)
    p_toc_end._element.append(r4)
    doc.add_page_break()

    # List of Figures
    p_lof_title = doc.add_paragraph()
    r_lof_t = p_lof_title.add_run("List of Figures")
    r_lof_t.bold = True
    r_lof_t.font.name = 'Book Antiqua'
    r_lof_t.font.size = Pt(14)
    p_lof_title.style = 'Heading 2'
    p_lof_title.paragraph_format.space_before = Pt(12)
    p_lof_title.paragraph_format.space_after = Pt(6)
    
    p_lof_begin = doc.add_paragraph()
    r_l1 = OxmlElement('w:r')
    fldChar_l1 = OxmlElement('w:fldChar')
    fldChar_l1.set(qn('w:fldCharType'), 'begin')
    r_l1.append(fldChar_l1)
    p_lof_begin._element.append(r_l1)
    
    r_l2 = OxmlElement('w:r')
    instrText_l2 = OxmlElement('w:instrText')
    instrText_l2.set(qn('xml:space'), 'preserve')
    instrText_l2.text = 'TOC \\h \\z \\t "Caption" \\c "Figure"'
    r_l2.append(instrText_l2)
    p_lof_begin._element.append(r_l2)
    
    r_l3 = OxmlElement('w:r')
    fldChar_l2 = OxmlElement('w:fldChar')
    fldChar_l2.set(qn('w:fldCharType'), 'separate')
    r_l3.append(fldChar_l2)
    p_lof_begin._element.append(r_l3)
    
    for fig_text in figures:
        p_entry = doc.add_paragraph(style='toc 2')
        p_entry.paragraph_format.space_before = Pt(0)
        p_entry.paragraph_format.space_after = Pt(2)
        r_txt = p_entry.add_run(fig_text)
        r_txt.font.name = 'Book Antiqua'
        r_txt.font.size = Pt(10)
        p_entry.add_run('\t')
        r_pg = p_entry.add_run('1')
        r_pg.font.name = 'Book Antiqua'
        r_pg.font.size = Pt(10)
        
    p_lof_end = doc.add_paragraph()
    r_l4 = OxmlElement('w:r')
    fldChar_l3 = OxmlElement('w:fldChar')
    fldChar_l3.set(qn('w:fldCharType'), 'end')
    r_l4.append(fldChar_l3)
    p_lof_end._element.append(r_l4)
    doc.add_page_break()

    # List of Tables
    p_lot_title = doc.add_paragraph()
    r_lot_t = p_lot_title.add_run("List of Tables")
    r_lot_t.bold = True
    r_lot_t.font.name = 'Book Antiqua'
    r_lot_t.font.size = Pt(14)
    p_lot_title.style = 'Heading 2'
    p_lot_title.paragraph_format.space_before = Pt(12)
    p_lot_title.paragraph_format.space_after = Pt(6)
    
    p_lot = doc.add_paragraph()
    add_field_code(p_lot, 'TOC \\h \\z \\t "Caption" \\c "Table"')
    doc.add_page_break()

    print("Step 7: Ingesting and parsing PROJECT_DOCUMENTATION.md...")
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    in_mermaid_block = False
    code_lines = []
    
    in_table = False
    table_rows = []
    
    fig_counter = 1
    current_heading = ""
    
    for line_idx, line in enumerate(lines):
        striped = line.strip()
        
        # --- CODE BLOCK HANDLING ---
        if striped.startswith("```"):
            if not in_code_block:
                in_code_block = True
                in_mermaid_block = striped.startswith("```mermaid")
                code_lines = []
            else:
                # End of code block
                if in_mermaid_block:
                    mermaid_text = "\n".join(code_lines)
                    heading_lower = current_heading.lower()
                    
                    if "2.1" in heading_lower:
                        fig_title = f"Figure {fig_counter}: High-Level Client-Server Architecture Diagram"
                        fig_counter += 1
                        
                        img_rendered = get_mermaid_image(mermaid_text, "rendered_high_level.png")
                        
                        p_cap = doc.add_paragraph()
                        p_cap.style = 'Caption'
                        p_cap.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        p_cap.paragraph_format.space_before = Pt(12)
                        p_cap.paragraph_format.space_after = Pt(6)
                        r_cap = p_cap.add_run(fig_title)
                        r_cap.bold = True
                        r_cap.font.name = 'Book Antiqua'
                        r_cap.font.size = Pt(10)
                        
                        p_img = doc.add_paragraph()
                        p_img.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        
                        if img_rendered and os.path.exists(img_rendered):
                            p_img.add_run().add_picture(img_rendered, width=Inches(6.2))
                        else:
                            print("    Fallback: Adding local image6.png and image7.png")
                            img_p6 = os.path.join(scratch_dir, 'image6.png')
                            if os.path.exists(img_p6):
                                p_img.add_run().add_picture(img_p6, width=Inches(2.5))
                                
                            p_img.add_run("      ")
                            
                            img_p7 = os.path.join(scratch_dir, 'image7.png')
                            if os.path.exists(img_p7):
                                p_img.add_run().add_picture(img_p7, width=Inches(3.8))
                            
                    elif "5.1" in heading_lower:
                        fig_title = f"Figure {fig_counter}: Entity Relationship Diagram (ERD) Schema"
                        fig_counter += 1
                        
                        img_rendered = get_mermaid_image(mermaid_text, "rendered_erd.png")
                        
                        p_cap = doc.add_paragraph()
                        p_cap.style = 'Caption'
                        p_cap.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        p_cap.paragraph_format.space_before = Pt(12)
                        p_cap.paragraph_format.space_after = Pt(6)
                        r_cap = p_cap.add_run(fig_title)
                        r_cap.bold = True
                        r_cap.font.name = 'Book Antiqua'
                        r_cap.font.size = Pt(10)
                        
                        p_img = doc.add_paragraph()
                        p_img.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        
                        if img_rendered and os.path.exists(img_rendered):
                            p_img.add_run().add_picture(img_rendered, width=Inches(6.2))
                        else:
                            print("    Fallback: Adding local image47.png")
                            img_p47 = os.path.join(scratch_dir, 'image47.png')
                            if os.path.exists(img_p47):
                                p_img.add_run().add_picture(img_p47, width=Inches(6.2))
                            
                    else:
                        fig_label = "Mermaid Sequence Diagram"
                        if "2.6.1" in heading_lower:
                            fig_label = "Dynamic Sales Excel Ingestion Sequence Flow"
                        elif "2.6.2" in heading_lower:
                            fig_label = "Agentic Chatbot RAG Query Resolution Process"
                        elif "2.6.3" in heading_lower:
                            fig_label = "User Authentication and RBAC Flow"
                        elif "2.6.4" in heading_lower:
                            fig_label = "Double-Entry Ledger Transaction Posting Flow"
                        elif "2.6.5" in heading_lower:
                            fig_label = "Supplier Procurement Ordered Slip Lifecycle"
                            
                        fig_title = f"Figure {fig_counter}: {fig_label}"
                        fig_counter += 1
                        
                        img_filename = f"rendered_seq_{fig_counter}.png"
                        img_rendered = get_mermaid_image(mermaid_text, img_filename)
                        
                        p_cap = doc.add_paragraph()
                        p_cap.style = 'Caption'
                        p_cap.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        p_cap.paragraph_format.space_before = Pt(12)
                        p_cap.paragraph_format.space_after = Pt(6)
                        r_cap = p_cap.add_run(fig_title)
                        r_cap.bold = True
                        r_cap.font.name = 'Book Antiqua'
                        r_cap.font.size = Pt(10)
                        
                        if img_rendered and os.path.exists(img_rendered):
                            p_img = doc.add_paragraph()
                            p_img.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                            p_img.add_run().add_picture(img_rendered, width=Inches(6.2))
                        else:
                            print(f"    Fallback: Adding Mermaid raw text for sequence diagram {fig_counter}")
                            p_code_header = doc.add_paragraph()
                            p_code_header.paragraph_format.space_before = Pt(6)
                            p_code_header.paragraph_format.space_after = Pt(2)
                            r_ch = p_code_header.add_run(f"[{fig_title} (Mermaid Source Code)]")
                            r_ch.bold = True
                            r_ch.font.name = 'Courier New'
                            r_ch.font.size = Pt(9.5)
                            
                            for line in code_lines:
                                p_c = doc.add_paragraph()
                                p_c.paragraph_format.left_indent = Inches(0.5)
                                p_c.paragraph_format.right_indent = Inches(0.5)
                                p_c.paragraph_format.space_after = Pt(0)
                                p_c.paragraph_format.line_spacing = 1.0
                                set_paragraph_shading(p_c, "F5F5F5")
                                r_c = p_c.add_run(line)
                                r_c.font.name = 'Courier New'
                                r_c.font.size = Pt(8.5)
                else:
                    # Regular Code block
                    for line in code_lines:
                        p_c = doc.add_paragraph()
                        p_c.paragraph_format.left_indent = Inches(0.5)
                        p_c.paragraph_format.right_indent = Inches(0.5)
                        p_c.paragraph_format.space_after = Pt(0)
                        p_c.paragraph_format.line_spacing = 1.0
                        set_paragraph_shading(p_c, "F8F9FA")
                        
                        r_c = p_c.add_run(line)
                        r_c.font.name = 'Courier New'
                        r_c.font.size = Pt(9)
                        r_c.font.color.rgb = RGBColor(36, 41, 47)
                        
                    p_sp = doc.add_paragraph()
                    p_sp.paragraph_format.space_after = Pt(6)
                    p_sp.paragraph_format.space_before = Pt(0)
                    p_sp.paragraph_format.line_spacing = 0.2
                    
                in_code_block = False
                in_mermaid_block = False
            continue
            
        if in_code_block:
            code_lines.append(line.replace('\n', ''))
            continue

        # --- TABLE INGESTION ---
        if striped.startswith("|"):
            cols = [c.strip() for c in line.split("|")[1:-1]]
            is_sep = all(re.match(r'^:?-+:?$', c) for c in cols)
            if not is_sep:
                table_rows.append(cols)
            in_table = True
            continue
        elif in_table:
            # End of table
            if table_rows:
                num_rows = len(table_rows)
                num_cols = len(table_rows[0])
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                
                for r_idx, row_vals in enumerate(table_rows):
                    for c_idx, val in enumerate(row_vals):
                        cell = table.cell(r_idx, c_idx)
                        cell.paragraphs[0].text = ""
                        p_cell = cell.paragraphs[0]
                        p_cell.paragraph_format.space_after = Pt(3)
                        p_cell.paragraph_format.space_before = Pt(3)
                        add_markdown_runs(p_cell, val)
                        if r_idx == 0:
                            for run in p_cell.runs:
                                run.bold = True
                
                doc.add_paragraph() # spacing
            table_rows = []
            in_table = False

        if not striped:
            continue
            
        # --- HEADINGS PARSING ---
        if striped.startswith("# "):
            title_text = striped[2:].strip()
            current_heading = title_text
            p_h = doc.add_paragraph()
            p_h.style = 'Title'
            p_h.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            p_h.paragraph_format.space_before = Pt(24)
            p_h.paragraph_format.space_after = Pt(12)
            r = p_h.add_run(title_text)
            r.bold = True
            r.font.name = 'Book Antiqua'
            r.font.size = Pt(20)
            
        elif striped.startswith("## "):
            h_text = striped[3:].strip()
            current_heading = h_text
            p_h = doc.add_paragraph(style='Heading 1')
            p_h.paragraph_format.space_before = Pt(18)
            p_h.paragraph_format.space_after = Pt(8)
            p_h.paragraph_format.keep_with_next = True
            add_markdown_runs(p_h, h_text)
            
        elif striped.startswith("### "):
            h_text = striped[4:].strip()
            current_heading = h_text
            p_h = doc.add_paragraph(style='Heading 2')
            p_h.paragraph_format.space_before = Pt(14)
            p_h.paragraph_format.space_after = Pt(6)
            p_h.paragraph_format.keep_with_next = True
            add_markdown_runs(p_h, h_text)
            
        elif striped.startswith("#### "):
            h_text = striped[5:].strip()
            current_heading = h_text
            p_h = doc.add_paragraph(style='Heading 3')
            p_h.paragraph_format.space_before = Pt(12)
            p_h.paragraph_format.space_after = Pt(6)
            p_h.paragraph_format.keep_with_next = True
            add_markdown_runs(p_h, h_text)
            
        # --- IMAGES PARSING ---
        elif striped.startswith("![") and striped.endswith(")"):
            match = re.match(r'^!\[(.*?)\]\((.*?)\)', striped)
            if match:
                caption_text = match.group(1).strip()
                img_url = match.group(2).strip()
                
                # Convert file:/// absolute path to local windows path
                local_path = img_url
                if local_path.startswith("file:///"):
                    local_path = local_path[8:]
                local_path = os.path.normpath(local_path)
                
                # Check if image exists
                if os.path.exists(local_path):
                    print(f"  Inserting UI screenshot: {local_path}")
                    p_img = doc.add_paragraph()
                    p_img.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(12)
                    p_img.paragraph_format.space_after = Pt(6)
                    p_img.add_run().add_picture(local_path, width=Inches(6.2))
                    
                    p_cap = doc.add_paragraph(style='Caption')
                    p_cap.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_before = Pt(6)
                    p_cap.paragraph_format.space_after = Pt(18)
                    
                    fig_title = f"Figure {fig_counter}: {caption_text}"
                    fig_counter += 1
                    
                    r_cap = p_cap.add_run(fig_title)
                    r_cap.bold = True
                    r_cap.font.name = 'Book Antiqua'
                    r_cap.font.size = Pt(10)
                else:
                    print(f"  Warning: image path not found: {local_path}")

        # --- LISTS PARSING ---
        elif striped.startswith("- ") or striped.startswith("* "):
            list_text = striped[2:].strip()
            p_l = doc.add_paragraph(style='List Paragraph')
            p_l.paragraph_format.space_before = Pt(0)
            p_l.paragraph_format.space_after = Pt(3)
            p_l.add_run('•\t')
            add_markdown_runs(p_l, list_text)
            
        elif re.match(r'^\d+\.\s+', striped):
            match = re.match(r'^(\d+\.)\s+(.*)', striped)
            num_prefix = match.group(1)
            list_text = match.group(2).strip()
            p_l = doc.add_paragraph(style='List Paragraph')
            p_l.paragraph_format.space_before = Pt(0)
            p_l.paragraph_format.space_after = Pt(3)
            p_l.add_run(f'{num_prefix}\t')
            add_markdown_runs(p_l, list_text)
            
        # --- STANDARD PARAGRAPHS ---
        else:
            p_p = doc.add_paragraph(style='Normal')
            p_p.paragraph_format.space_before = Pt(0)
            p_p.paragraph_format.space_after = Pt(6)
            p_p.paragraph_format.line_spacing = 1.15
            add_markdown_runs(p_p, striped)

    print("\nStep 8: Updating headers and footers project name...")
    for sec_idx, section in enumerate(doc.sections):
        # Footers
        for p in section.footer.paragraphs:
            for r in p.runs:
                if '<Project Name>' in r.text:
                    r.text = r.text.replace('<Project Name>', 'Bizionary')
        # Headers
        for p in section.header.paragraphs:
            for r in p.runs:
                if '<Project Name>' in r.text:
                    r.text = r.text.replace('<Project Name>', 'Bizionary')

    print("\nStep 9: Saving reformatted document...")
    doc.save(output_path)
    print(f"Success! Reformatted document saved to:\n  {output_path}")

if __name__ == "__main__":
    main()
